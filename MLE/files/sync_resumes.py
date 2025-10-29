#!/usr/bin/env python3
"""
Resume Synchronization Script
Syncs multiple resume Word documents that differ only by email address,
then converts them all to PDF.
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
import shutil

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it by running:")
    print("  pip3 install python-docx")
    print("\nOr if that doesn't work, try:")
    print("  pip3 install --user python-docx")
    sys.exit(1)


# Define your email addresses
EMAILS = [
    "niketj@cs.cmu.edu",
    "nikj1301@gmail.com",
    "niketj@andrew.cmu.edu",
    "niketjaina4@gmail.com",
    "niketjainapp@gmail.com"
]

# PDF output filename (all PDFs will be named this)
PDF_OUTPUT_NAME = "Resume_Niket_Jain.pdf"


def find_email_in_document(doc):
    """Find email address in a Word document."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        emails = re.findall(email_pattern, paragraph.text)
        if emails:
            return emails[0]
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                emails = re.findall(email_pattern, cell.text)
                if emails:
                    return emails[0]
    
    # Search in headers and footers
    for section in doc.sections:
        # Check header
        for paragraph in section.header.paragraphs:
            emails = re.findall(email_pattern, paragraph.text)
            if emails:
                return emails[0]
        
        # Check footer
        for paragraph in section.footer.paragraphs:
            emails = re.findall(email_pattern, paragraph.text)
            if emails:
                return emails[0]
    
    return None


def replace_email_in_document(doc, old_email, new_email):
    """Replace email address throughout the document."""
    
    def replace_in_paragraph(paragraph):
        """Replace email in a paragraph, handling text split across runs."""
        # First, check if email exists in the full paragraph text
        if old_email not in paragraph.text:
            return False
        
        # Get the full text and find email position
        full_text = paragraph.text
        email_start = full_text.find(old_email)
        
        if email_start == -1:
            return False
        
        # Build a map of character positions to runs
        char_to_run = []
        current_pos = 0
        
        for run in paragraph.runs:
            run_len = len(run.text)
            for i in range(run_len):
                char_to_run.append(run)
            current_pos += run_len
        
        # Find which runs contain the email
        email_end = email_start + len(old_email)
        affected_runs = set()
        
        for pos in range(email_start, min(email_end, len(char_to_run))):
            if pos < len(char_to_run):
                affected_runs.add(char_to_run[pos])
        
        if not affected_runs:
            return False
        
        # Strategy: Replace text in all affected runs
        # Clear all affected runs first
        for run in affected_runs:
            run.text = ""
        
        # Put the new email in the first affected run
        first_run = list(affected_runs)[0]
        
        # Reconstruct the paragraph text with the replacement
        new_text = full_text[:email_start] + new_email + full_text[email_end:]
        
        # Clear all runs and set the text in the first run
        for run in paragraph.runs:
            run.text = ""
        
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        
        return True
    
    replaced = False
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph):
            replaced = True
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph):
                        replaced = True
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            if replace_in_paragraph(paragraph):
                replaced = True
        
        # Footer
        for paragraph in section.footer.paragraphs:
            if replace_in_paragraph(paragraph):
                replaced = True
    
    return replaced


def find_all_resume_files(base_dir):
    """Find all .docx resume files in subdirectories."""
    resume_files = []
    base_path = Path(base_dir)
    
    for docx_file in base_path.rglob("*.docx"):
        # Skip temporary Word files
        if not docx_file.name.startswith("~$"):
            resume_files.append(docx_file)
    
    return resume_files


def get_most_recent_file(files):
    """Get the most recently modified file from a list."""
    if not files:
        return None
    
    return max(files, key=lambda f: f.stat().st_mtime)


def convert_docx_to_pdf(docx_path, output_name="Resume_Niket_Jain.pdf"):
    """Convert a .docx file to PDF using LibreOffice."""
    import subprocess
    
    docx_path = Path(docx_path)
    output_dir = docx_path.parent
    
    # Try different LibreOffice command locations (macOS, Linux, Windows)
    libreoffice_commands = [
        'soffice',  # Common on macOS and Linux
        'libreoffice',  # Linux
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS default
        '/usr/bin/soffice',  # Linux alternative
        '/usr/bin/libreoffice',  # Linux alternative
    ]
    
    for cmd in libreoffice_commands:
        try:
            result = subprocess.run([
                cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(docx_path)
            ], check=True, capture_output=True, text=True, timeout=30)
            
            # LibreOffice creates PDF with same name as docx
            temp_pdf = output_dir / f"{docx_path.stem}.pdf"
            final_pdf = output_dir / output_name
            
            # Rename to desired output name (replacing if exists)
            if temp_pdf.exists():
                if final_pdf.exists():
                    final_pdf.unlink()  # Remove existing file
                temp_pdf.rename(final_pdf)
                print(f"  ✓ Converted to PDF: {output_name}")
            else:
                print(f"  ✗ Error: PDF file was not created")
                return False
            
            return True
        except subprocess.CalledProcessError as e:
            # Command found but conversion failed
            print(f"  ✗ Error converting {docx_path.name} to PDF:")
            if e.stderr:
                print(f"     {e.stderr.strip()}")
            return False
        except FileNotFoundError:
            # This command doesn't exist, try the next one
            continue
        except subprocess.TimeoutExpired:
            print(f"  ✗ Error: Conversion timed out for {docx_path.name}")
            return False
    
    # If we get here, none of the commands worked
    print("  ✗ Error: LibreOffice not found. Please install it.")
    print("    On macOS: brew install --cask libreoffice")
    print("    On Linux: sudo apt-get install libreoffice")
    return False


def sync_resumes(base_dir):
    """Main function to sync all resumes."""
    print("=" * 60)
    print("Resume Synchronization Script")
    print("=" * 60)
    
    # Find all resume files
    resume_files = find_all_resume_files(base_dir)
    
    if not resume_files:
        print(f"Error: No .docx files found in {base_dir}")
        return
    
    print(f"\nFound {len(resume_files)} resume file(s):")
    for f in resume_files:
        mod_time = datetime.fromtimestamp(f.stat().st_mtime)
        print(f"  - {f.name} (modified: {mod_time.strftime('%Y-%m-%d %H:%M:%S')})")
    
    # Find the most recently modified file
    source_file = get_most_recent_file(resume_files)
    print(f"\n📄 Source (most recent): {source_file.name}")
    
    # Load the source document
    try:
        source_doc = Document(source_file)
    except Exception as e:
        print(f"Error loading source document: {e}")
        return
    
    # Find email in source document
    source_email = find_email_in_document(source_doc)
    if not source_email:
        print("Warning: No email found in source document")
        source_email = None
    else:
        print(f"📧 Source email: {source_email}")
    
    print(f"\n{'=' * 60}")
    print("Updating resumes...")
    print("=" * 60)
    
    # Update all other resume files
    for target_file in resume_files:
        if target_file == source_file:
            print(f"\n✓ {target_file.name} (source - skipped)")
            continue
        
        print(f"\n📝 Updating: {target_file.name}")
        
        try:
            # Load target document to get its email
            target_doc = Document(target_file)
            target_email = find_email_in_document(target_doc)
            
            if not target_email:
                print(f"  Warning: No email found in {target_file.name}")
                target_email = None
            else:
                print(f"  Current email: {target_email}")
            
            # Create a copy of the source document
            temp_file = target_file.with_suffix('.docx.tmp')
            shutil.copy2(source_file, temp_file)
            
            # Load the temp copy
            new_doc = Document(temp_file)
            
            # Replace email if both emails exist
            if source_email and target_email:
                replaced = replace_email_in_document(new_doc, source_email, target_email)
                if replaced:
                    print(f"  ✓ Replaced {source_email} → {target_email}")
                else:
                    print(f"  ⚠ Warning: Could not replace email (email may be formatted in a complex way)")
            elif target_email:
                print(f"  ℹ Keeping target email: {target_email}")
            else:
                print(f"  ℹ No email found to replace")
            
            # Save the updated document
            new_doc.save(target_file)
            temp_file.unlink()  # Remove temp file
            
            # Verify the email was replaced
            verify_doc = Document(target_file)
            verify_email = find_email_in_document(verify_doc)
            
            if verify_email == target_email:
                print(f"  ✓ Updated successfully - Email verified: {verify_email}")
            elif verify_email:
                print(f"  ⚠ Warning: Email in saved file is {verify_email}, expected {target_email}")
            else:
                print(f"  ⚠ Warning: No email found in saved file")
            
        except Exception as e:
            print(f"  ✗ Error updating {target_file.name}: {e}")
    
    print(f"\n{'=' * 60}")
    print("Converting to PDF...")
    print("=" * 60)
    
    # Convert all docx files to PDF
    for docx_file in resume_files:
        print(f"\n📄 {docx_file.name}")
        convert_docx_to_pdf(docx_file, PDF_OUTPUT_NAME)
    
    print(f"\n{'=' * 60}")
    print("✓ Synchronization complete!")
    print("=" * 60)


if __name__ == "__main__":
    if len(sys.argv) > 1:
        base_directory = sys.argv[1]
    else:
        # Use current directory if no argument provided
        base_directory = "."
    
    if not os.path.exists(base_directory):
        print(f"Error: Directory '{base_directory}' does not exist")
        sys.exit(1)
    
    sync_resumes(base_directory)